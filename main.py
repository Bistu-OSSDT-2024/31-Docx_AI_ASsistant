from zhipuai import ZhipuAI
from docx import Document
import tkinter as tk
from tkinter import ttk  # 导入 ttk 模块用于更美观的控件
from tkinter import filedialog
from ttkthemes import ThemedTk

def correction(content):
    client = ZhipuAI(api_key="73f54cc77008aa76975c51b196a1ab2e.80TZQ64Msp8ujxGQ")
    messages = [
        {"role": "user",
         "content": "作为一个纠错专家，请你找出这段文本中的错误，不要输出其他任何东西，直接输出纠错后的结果"},
        {"role": "assistant", "content": "当然，请给出你的文本"},
        {"role": "user", "content": content.text}
    ]
    response = client.chat.completions.create(
        model="glm-4",
        messages=messages,
        tool_choice="auto",
    )
    return response.choices[0].message.content

#总结
def summarize_from_docx(content):
    client = ZhipuAI(api_key="73f54cc77008aa76975c51b196a1ab2e.80TZQ64Msp8ujxGQ")
    messages = [
        {"role": "user",
         "content": "为下面的文档写出总结，我将分段发送给你这个文章，我可能会分段发送，不要输出任何别的东西,直接输出结果"},
        {"role": "assistant", "content": "当然，请给出你的文本"},
        {"role": "user", "content": content.text}
    ]
    response = client.chat.completions.create(
        model="glm-4",
        messages=messages,
        tool_choice="auto",
    )
    return response.choices[0].message.content

def main(input_path, output_path, mode):
    doc = Document(input_path)
    corrected_doc = Document()
    if mode == '纠错':
        for para in doc.paragraphs:
            corrected_doc.add_paragraph().add_run(correction(para))
    elif mode == '总结':
        for para in doc.paragraphs:
            corrected_doc.add_paragraph().add_run(summarize_from_docx(para))
    corrected_doc.save(output_path)
    print(f"{mode}完成，结果保存在 {output_path}")
    root.destroy()  # 在处理完成后关闭窗口

# ui
root = ThemedTk(theme="arc")  # 选择一个主题，这里是 'arc'，您可以选择其他喜欢的主题

def select_input_file():
    input_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    input_entry.delete(0, tk.END)
    input_entry.insert(0, input_path)

# 选择输出文件的函数
def select_output_file():
    output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
    output_entry.delete(0, tk.END)
    output_entry.insert(0, output_path)

# 开始处理的函数
def start_process():
    input_path = input_entry.get()
    output_path = output_entry.get()
    mode = mode_var.get()  # 获取选择的模式
    main(input_path, output_path, mode)

# 创建输入文件标签
input_label = ttk.Label(root, text="输入文件:", font=("Helvetica", 12))  # 使用 ttk.Label 并设置字体
input_label.grid(row=0, column=0, padx=10, pady=10)  # 增加内边距

# 创建输入文件输入框
input_entry = ttk.Entry(root, width=30)  # 设置宽度
input_entry.grid(row=0, column=1, padx=10, pady=10)

# 创建输入文件选择按钮
input_button = ttk.Button(root, text="选择", command=select_input_file)
input_button.grid(row=0, column=2, padx=10, pady=10)

# 创建输出文件标签
output_label = ttk.Label(root, text="输出文件:", font=("Helvetica", 12))
output_label.grid(row=1, column=0, padx=10, pady=10)

# 创建输出文件输入框
output_entry = ttk.Entry(root, width=30)
output_entry.grid(row=1, column=1, padx=10, pady=10)

# 创建输出文件选择按钮
output_button = ttk.Button(root, text="选择", command=select_output_file)
output_button.grid(row=1, column=2, padx=10, pady=10)

# 创建模式选择标签
mode_label = ttk.Label(root, text="选择模式:", font=("Helvetica", 12))
mode_label.grid(row=2, column=0, padx=10, pady=10)

# 创建模式选择变量
mode_var = tk.StringVar()
mode_var.set('纠错')  # 默认选择纠错

# 创建纠错选项
correction_option = ttk.Radiobutton(root, text="纠错", variable=mode_var, value='纠错')
correction_option.grid(row=2, column=1, padx=10, pady=10)

# 创建总结选项
summarize_option = ttk.Radiobutton(root, text="总结", variable=mode_var, value='总结')
summarize_option.grid(row=2, column=2, padx=10, pady=10)

# 创建开始处理按钮
start_button = ttk.Button(root, text="开始处理", command=start_process)
start_button.grid(row=3, column=1, padx=10, pady=20)

root.mainloop()