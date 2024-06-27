from zhipuai import ZhipuAI
from docx import Document
import tkinter as tk
from tkinter import ttk, messagebox  # 导入 ttk 模块用于更美观的控件
from tkinter import filedialog
from ttkthemes import ThemedTk


def correction(content):
    client = ZhipuAI(api_key="73f54cc77008aa76975c51b196a1ab2e.80TZQ64Msp8ujxGQ")
    messages = [
        {"role": "user",
         "content": "作为一个纠错专家，请你找出这段文本中的错误，不要输出其他任何东西，直接输出纠错后的结果"},        # 1.2 改进大模型训练
        {"role": "assistant", "content": "当然，请给出你的文本"},
        {"role": "user", "content": content}
    ]
    response = client.chat.completions.create(
        model="glm-4",
        messages=messages,
        tool_choice="auto",
    )
    return response.choices[0].message.content


# 总结
def summarize_from_docx(content):
    client = ZhipuAI(api_key="73f54cc77008aa76975c51b196a1ab2e.80TZQ64Msp8ujxGQ")
    messages = [
        {"role": "user",
         "content": "为下面的文档写出总结，不要输出任何别的东西,直接输出结果"},    # 1.2 改进大模型训练
        {"role": "assistant", "content": "当然，请给出你的文本"},
        {"role": "user", "content": content}
    ]
    response = client.chat.completions.create(
        model="glm-4",
        messages=messages,
        tool_choice="auto",
    )
    return response.choices[0].message.content
# 改进建议
def summarize2_from_docx(content):
    client = ZhipuAI(api_key="73f54cc77008aa76975c51b196a1ab2e.80TZQ64Msp8ujxGQ")
    messages = [
        {"role": "user",
         "content": "为下面的文档写出改进建议"},
        {"role": "assistant", "content": "当然，请给出你的文档"},
        {"role": "user", "content": content}
    ]
    response = client.chat.completions.create(
        model="glm-4",
        messages=messages,
        tool_choice="auto",
    )
    return response.choices[0].message.content

# v1.2 增加改进建议
# v1.1 计算字数
def count_words_in_paragraph(paragraph):
    # 去除段落中的空格和换行符
    text = paragraph
    # 返回文本的长度，即字数
    return len(text)

# 主函数
# v1.0进行分段传输
# v1.1增加限制，空格段不再传输并且删除的分段传输
def main(input_path, output_path, mode):
    doc = Document(input_path)
    corrected_doc = Document()
    if mode == '纠错':
        combined_paragraph = ""  # 用于存储组合的段落内容
        for para in doc.paragraphs:
            if para != "":
                combined_paragraph += para.text + '\n'  # 1.1 将段落对象转换为文本后累加，并添加换行符
                if count_words_in_paragraph(combined_paragraph) > 3000:  # 1.1 达到 3000 字及以上
                    corrected_doc.add_paragraph().add_run(correction(combined_paragraph))
                    combined_paragraph = ""  # 重置
        corrected_doc.add_paragraph().add_run(correction(combined_paragraph))   # 执行剩下的
    elif mode == '总结':
        combined_paragraph = ""
        for para in doc.paragraphs:
            if para != "":
                combined_paragraph += para.text + '\n'  # 1.1 将段落对象转换为文本后累加，并添加换行符
                if count_words_in_paragraph(combined_paragraph) > 3000:  # 1.1 达到 3000 字及以上
                    corrected_doc.add_paragraph().add_run(summarize_from_docx(combined_paragraph))
                    combined_paragraph = ""  # 重置
        corrected_doc.add_paragraph().add_run(summarize_from_docx(combined_paragraph))  # 执行剩下的
    elif mode == '改进建议':
        combined_paragraph = ""
        for para in doc.paragraphs:
            if para != "":
                combined_paragraph += para.text + '\n'
                if count_words_in_paragraph(combined_paragraph) > 3000:
                    result = summarize2_from_docx(combined_paragraph)
                    show_improvement_suggestion(result)  # 调用新的函数来显示结果
                    combined_paragraph = ""
        result = summarize2_from_docx(combined_paragraph)
        show_improvement_suggestion(result)  # 显示剩余部分的结果
    corrected_doc.save(output_path)
    print(f"{mode}完成，结果保存在 {output_path}")
    messagebox.showinfo("提示", "操作完成，可以继续使用")

def show_improvement_suggestion(result):
    # 创建一个新的窗口来显示改进建议
    improvement_window = tk.Toplevel(root)
    improvement_window.title("改进建议")

    # 创建一个文本框来显示结果
    text_box = tk.Text(improvement_window, wrap=tk.WORD, width=50, height=20)
    text_box.pack()
    text_box.insert(tk.END, result)

    # 创建确定按钮
    ok_button = ttk.Button(improvement_window, text="确定", command=improvement_window.destroy)
    ok_button.pack()

# ui v1.2
root = ThemedTk(theme="arc")


def select_input_file():
    input_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    input_entry.delete(0, tk.END)
    input_entry.insert(0, input_path)


# 选择输出文件的函数
def select_output_file():
    output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
    output_entry.delete(0, tk.END)
    output_entry.insert(0, output_path)


# 开始处理的函数
def start_process():
    input_path = input_entry.get()
    output_path = output_entry.get()
    mode = mode_var.get()  # 获取选择的模式
    main(input_path, output_path, mode)


# 创建输入文件标签
input_label = ttk.Label(root, text="输入文件:", font=("Helvetica", 12))  # 使用 ttk.Label 并设置字体
input_label.grid(row=0, column=0, padx=10, pady=10)  # 增加内边距

# 创建输入文件输入框
input_entry = ttk.Entry(root, width=30)  # 设置宽度
input_entry.grid(row=0, column=1, padx=10, pady=10)

# 创建输入文件选择按钮
input_button = ttk.Button(root, text="选择", command=select_input_file)
input_button.grid(row=0, column=2, padx=10, pady=10)

# 创建输出文件标签
output_label = ttk.Label(root, text="输出文件:", font=("Helvetica", 12))
output_label.grid(row=1, column=0, padx=10, pady=10)

# 创建输出文件输入框
output_entry = ttk.Entry(root, width=30)
output_entry.grid(row=1, column=1, padx=10, pady=10)

# 创建输出文件选择按钮
output_button = ttk.Button(root, text="选择", command=select_output_file)
output_button.grid(row=1, column=2, padx=10, pady=10)

# 创建模式选择标签
mode_label = ttk.Label(root, text="选择模式:", font=("Helvetica", 12))
mode_label.grid(row=2, column=0, padx=10, pady=10)

# 创建模式选择变量
mode_var = tk.StringVar()
mode_var.set('纠错')  # 默认选择纠错

# 创建纠错选项
correction_option = ttk.Radiobutton(root, text="纠错", variable=mode_var, value='纠错')
correction_option.grid(row=2, column=1, padx=10, pady=10)

# 创建总结选项
summarize_option = ttk.Radiobutton(root, text="总结", variable=mode_var, value='总结')
summarize_option.grid(row=2, column=2, padx=10, pady=10)

# 创建改进建议选项
improvement_option = ttk.Radiobutton(root, text="改进建议", variable=mode_var, value='改进建议')
improvement_option.grid(row=2, column=3, padx=10, pady=10)

# 创建开始处理按钮
start_button = ttk.Button(root, text="开始处理", command=start_process)
start_button.grid(row=3, column=1, padx=10, pady=20)

root.mainloop()
